"""Executable conformance audit for the uploaded NFL FIDOS Master Plan.

The audit is intentionally evidence-oriented.  It does not declare the program
complete; it proves that the two source representations are structurally
reviewable and that the implementation ledger points to real repository
artifacts for each stage.
"""

from __future__ import annotations

import hashlib
import json
import re
import zipfile
from xml.etree import ElementTree
from pathlib import Path
from typing import Any


STAGE_HEADING = re.compile(r"^##\s+STAGE\s+(\d+)\s+[—-]\s*(.+?)\s*$", re.IGNORECASE)
REQUIRED_HEADING = re.compile(r"^###\s+Required Deliverables\s*$", re.IGNORECASE)
NEXT_HEADING = re.compile(r"^#{1,3}\s+")
BULLET = re.compile(r"^\s*[-*]\s+(.+?)\s*$")
WORK_PACKAGE_ROW = re.compile(r"^\s*\|\s*([^|]+)\s*\|\s*([^|]+)\s*\|\s*([^|]+)\s*\|\s*([^|]+)\s*\|\s*$")


def _normalize(value: str) -> str:
    value = value.replace("â€”", "-").replace("—", "-").replace("–", "-")
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def parse_markdown_plan(path: str | Path) -> dict[str, Any]:
    """Extract all roadmap stages and their required deliverables."""
    source = Path(path)
    lines = source.read_text(encoding="utf-8").splitlines()
    stages: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    collecting = False
    collecting_work_packages = False
    for line_number, line in enumerate(lines, start=1):
        match = STAGE_HEADING.match(line.strip())
        if match:
            if current is not None:
                stages.append(current)
            current = {
                "stage": f"STAGE-{int(match.group(1))}",
                "number": int(match.group(1)),
                "title": match.group(2).strip(),
                "heading_line": line_number,
                "required_deliverables": [],
            }
            collecting = False
            collecting_work_packages = current["stage"] == "STAGE-0"
            continue
        if current is None:
            continue
        if REQUIRED_HEADING.match(line.strip()):
            collecting = True
            collecting_work_packages = False
            continue
        if collecting_work_packages:
            row = WORK_PACKAGE_ROW.match(line)
            if row and row.group(1).strip().lower() not in {"work package", "---"}:
                current["required_deliverables"].append(row.group(4).strip())
            elif line.strip().startswith("#"):
                collecting_work_packages = False
        if collecting and NEXT_HEADING.match(line.strip()):
            collecting = False
        if collecting:
            bullet = BULLET.match(line)
            if bullet:
                current["required_deliverables"].append(bullet.group(1).strip())
    if current is not None:
        stages.append(current)
    return {
        "path": str(source),
        "sha256": _sha256(source),
        "line_count": len(lines),
        "stages": stages,
    }


def extract_docx_text(path: str | Path) -> list[str]:
    """Extract paragraph and table text without requiring a rendering engine."""
    try:
        from docx import Document
    except ImportError:  # pragma: no cover - exercised when the optional helper is absent
        # A .docx is an OOXML package.  The conformance audit only needs text,
        # so keep it runnable in the minimal system interpreter as well as the
        # richer workspace runtime that provides python-docx.
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        try:
            with zipfile.ZipFile(path) as archive:
                root = ElementTree.fromstring(archive.read("word/document.xml"))
        except (OSError, KeyError, zipfile.BadZipFile, ElementTree.ParseError) as exc:
            raise RuntimeError("DOCX OOXML document.xml could not be read") from exc
        values: list[str] = []
        for paragraph in root.iter(namespace + "p"):
            text = "".join(node.text or "" for node in paragraph.iter(namespace + "t")).strip()
            if text:
                values.append(text)
        return values
    document = Document(str(path))
    values = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text.strip() for cell in row.cells if cell.text.strip())
    return values


def _docx_stage_matches(text: list[str], stage: dict[str, Any]) -> bool:
    number = str(stage["number"])
    title = _normalize(stage["title"])
    for item in text:
        normalized = _normalize(item)
        if re.search(rf"\bstage\s+{re.escape(number)}\b", normalized) and title[:24] in normalized:
            return True
    return False


def _evidence_exists(root: Path, evidence: str) -> bool:
    if evidence.startswith("EVAL-FAM-"):
        return True
    candidate = root / evidence
    return candidate.exists()


def audit_master_plan(
    markdown_path: str | Path,
    docx_path: str | Path,
    repository_root: str | Path,
    traceability_path: str | Path,
) -> dict[str, Any]:
    """Audit source parity, roadmap completeness, and evidence reachability."""
    root = Path(repository_root)
    markdown = parse_markdown_plan(markdown_path)
    docx = Path(docx_path)
    docx_text = extract_docx_text(docx)
    traceability = json.loads(Path(traceability_path).read_text(encoding="utf-8"))
    ledger_stages = {record.get("stage"): record for record in traceability.get("stages", [])}
    errors: list[str] = []
    warnings: list[str] = []
    stages = markdown["stages"]
    expected = [f"STAGE-{number}" for number in range(26)]
    actual = [stage["stage"] for stage in stages]
    if actual != expected:
        errors.append(f"Markdown roadmap must contain STAGE-0 through STAGE-25 in order; found {actual}")
    missing_docx_stages = [stage["stage"] for stage in stages if not _docx_stage_matches(docx_text, stage)]
    if missing_docx_stages:
        errors.append(f"DOCX is missing stage headings: {missing_docx_stages}")
    missing_deliverables = [stage["stage"] for stage in stages if not stage["required_deliverables"]]
    if missing_deliverables:
        errors.append(f"Stages without required deliverables: {missing_deliverables}")
    missing_ledger = [stage["stage"] for stage in stages if stage["stage"] not in ledger_stages]
    if missing_ledger:
        errors.append(f"Traceability ledger is missing stages: {missing_ledger}")
    unreachable: dict[str, list[str]] = {}
    for stage in stages:
        record = ledger_stages.get(stage["stage"], {})
        missing = [evidence for evidence in record.get("evidence", []) if not _evidence_exists(root, evidence)]
        if missing:
            unreachable[stage["stage"]] = missing
    if unreachable:
        errors.append("Traceability evidence paths do not resolve")
    for stage in stages:
        if len(stage["required_deliverables"]) < 1:
            continue
        if not ledger_stages.get(stage["stage"], {}).get("remaining"):
            warnings.append(f"{stage['stage']} has no explicit remaining-work statement")
    return {
        "audit_id": "AUDIT-MASTER-PLAN-CONFORMANCE-001",
        "status": "passed" if not errors else "failed",
        "markdown": {k: markdown[k] for k in ("path", "sha256", "line_count")},
        "docx": {"path": str(docx), "sha256": _sha256(docx), "paragraph_and_cell_count": len(docx_text)},
        "stage_count": len(stages),
        "required_deliverable_count": sum(len(stage["required_deliverables"]) for stage in stages),
        "docx_stage_count": sum(1 for stage in stages if stage["stage"] not in missing_docx_stages),
        "traceability_stage_count": len(ledger_stages),
        "unreachable_evidence": unreachable,
        "warnings": warnings,
        "errors": errors,
    }
